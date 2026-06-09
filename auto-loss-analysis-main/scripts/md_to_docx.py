#!/usr/bin/env python3
"""Convert the final report.md → report.docx with NVIDIA/Mercor styling.

Delegates to the canonical styled converter at:
  ~/.claude/skills/report-style-enforcer/styled_md_to_docx.py

Style guide reference:
  ~/.claude/skills/report-style-enforcer/STYLE_GUIDE.md

Extracted from: [EXT][Nvidia] Agentic Code Master Doc.docx (2026-04-14)

Formatting spec:
- H1: 20pt bold black Arial
- H2: 13.5pt bold black + bottom border (section dividers)
- H3: 13pt bold purple (#674EA7)
- Body: 10pt Arial
- Inline code: 10pt Roboto Mono green (#188038)
- Tables: purple header shading (#B4A7D6), borderless data tables, bordered comparison tables
- Key headlines: 0.5" indent, bold lead sentence
- Callout blockquotes: 0.3" indent, purple left border, italic
- Page: 8.5x11, 1" margins

Usage: python3 scripts/md_to_docx.py --input output/report.md --output output/report.docx
"""

import argparse
import os
import sys

# Try to import from the canonical skill location first
skill_path = os.path.expanduser("~/.claude/skills/report-style-enforcer")
if os.path.exists(os.path.join(skill_path, "styled_md_to_docx.py")):
    sys.path.insert(0, skill_path)
    from styled_md_to_docx import convert
else:
    # Inline fallback — same implementation
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    import re

    PURPLE_PRIMARY = RGBColor(0x67, 0x4E, 0xA7)
    PURPLE_LIGHT = "B4A7D6"
    PURPLE_LIGHTER = "D9D2E9"
    GREEN_CODE = RGBColor(0x18, 0x80, 0x38)
    BLACK = RGBColor(0x00, 0x00, 0x00)
    BODY_FONT = "Arial"
    CODE_FONT = "Roboto Mono"

    def convert(input_path, output_path):
        """Fallback converter — install the full skill for best results."""
        print("WARNING: Using fallback converter. Install ~/.claude/skills/report-style-enforcer/ for full styling.")
        # Minimal conversion
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = BODY_FONT
        style.font.size = Pt(10)
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        with open(input_path) as f:
            for line in f:
                line = line.rstrip()
                if line.startswith("# ") and not line.startswith("## "):
                    p = doc.add_paragraph()
                    run = p.add_run(line[2:])
                    run.font.size = Pt(20)
                    run.font.bold = True
                elif line.startswith("## "):
                    p = doc.add_paragraph()
                    run = p.add_run(line[3:])
                    run.font.size = Pt(13.5)
                    run.font.bold = True
                elif line.startswith("### "):
                    p = doc.add_paragraph()
                    run = p.add_run(line[4:])
                    run.font.size = Pt(13)
                    run.font.bold = True
                    run.font.color.rgb = PURPLE_PRIMARY
                elif line.strip():
                    doc.add_paragraph(line.strip())

        doc.save(output_path)
        print(f"Written to {output_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", default="./output/report.md")
    parser.add_argument("--output", default="./output/report.docx")
    args = parser.parse_args()
    convert(args.input, args.output)
